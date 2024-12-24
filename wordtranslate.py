import streamlit as st
import os
import google.generativeai as genai
from docx import Document
from io import BytesIO
from docx.shared import Inches

# Configure the GenAI client
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

def translate_text_gemini(text, source_language="en", target_language="ja"):
    """
    Translates text using Gemini 1.5 Pro.
    """
    try:
        generation_config = {
            "temperature": 1.0,
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 8192,
            "response_mime_type": "text/plain",
        }
        
        model = genai.GenerativeModel(model_name="gemini-1.5-pro", generation_config=generation_config)
        prompt = f"Translate the following text from {source_language} to {target_language}:\n{text}"
        response = model.generate_content(prompt)
        return response.text if response else text
    except Exception as e:
        st.error(f"Translation error: {e}")
        return text

def translate_document(file, source_language="en", target_language="ja"):
    """
    Reads and translates a Word document, preserving images.
    """
    doc = Document(file)
    translated_doc = Document()

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Translate non-empty paragraphs
            translated_text = translate_text_gemini(paragraph.text, source_language, target_language)
            translated_doc.add_paragraph(translated_text)
        else:
            translated_doc.add_paragraph("")  # Preserve empty paragraphs

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_stream = rel.target_part.blob
            new_paragraph = translated_doc.add_paragraph()
            new_paragraph.add_run().add_picture(BytesIO(image_stream), width=Inches(1.0))

    # Save translated document to BytesIO
    output_stream = BytesIO()
    translated_doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

def main():
    st.title("Document Translator with GenAI")
    
    uploaded_file = st.file_uploader("Upload a Word document (.docx) to translate")
    
    if uploaded_file is not None:
        st.write("Original document content:")
        
        doc = Document(uploaded_file)
        for paragraph in doc.paragraphs:
            st.write(paragraph.text)

        # Translate document
        st.write("Translating...")
        translated_file = translate_document(uploaded_file)
        
        # Downloadable translated document
        st.download_button(
            label="Download Translated Document",
            data=translated_file,
            file_name="translated_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
